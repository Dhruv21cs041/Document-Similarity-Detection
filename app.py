import streamlit as st
import spacy
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import PyPDF2
from docx import Document

# Load spaCy English model
nlp = spacy.load('en_core_web_sm')

def preprocess_text(text):
    # Tokenize and lemmatize text
    doc = nlp(text)
    tokens = [token.lemma_ for token in doc if not token.is_stop and token.is_alpha]
    return ' '.join(tokens)

def read_pdf(file_path):
    with open(file_path, 'rb') as file:
        pdf_reader = PyPDF2.PdfFileReader(file)
        text = ''
        for page_num in range(pdf_reader.numPages):
            text += pdf_reader.getPage(page_num).extractText()
        return text

def read_word(file_path):
    doc = Document(file_path)
    text = ''
    for paragraph in doc.paragraphs:
        text += paragraph.text + '\n'
    return text

def read_text(file):
    if isinstance(file, str):
        # If the input is a regular text, return as is
        return file
    elif hasattr(file, "name") and file.name.lower().endswith('.pdf'):
        # If the file is a PDF, use the PDF reader
        return read_pdf(file)
    elif hasattr(file, "name") and file.name.lower().endswith('.docx'):
        # If the file is a Word document, use the Word reader
        return read_word(file)
    else:
        # Unsupported file format
        raise ValueError("Unsupported file format")

def read_pdf(file):
    # Read the content of the PDF file
    text = ""
    with file as f:
        pdf_reader = PyPDF2.PdfReader(f)
        for page_num in range(len(pdf_reader.pages)):
            text += pdf_reader.pages[page_num].extract_text()
    return text

def read_word(file):
    # Read the content of the Word document
    text = ""
    with file as f:
        doc = Document(f)
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
    return text

def vectorize_documents(documents):
    tfidf_vectorizer = TfidfVectorizer(stop_words='english', preprocessor=preprocess_text)
    tfidf_matrix = tfidf_vectorizer.fit_transform(documents)
    return tfidf_matrix

def calculate_similarity(matrix):
    similarity_matrix = cosine_similarity(matrix, matrix)
    return similarity_matrix

def main():
    st.title('Document Similarity Detection')

    # Upload files
    uploaded_files = st.file_uploader('Upload documents (PDF or Word)', type=['pdf', 'docx'], accept_multiple_files=True)

    if uploaded_files:
        input_documents = [read_text(file) for file in uploaded_files]
        file_names = [file.name for file in uploaded_files]

        # Preprocess and vectorize documents
        tfidf_matrix = vectorize_documents(input_documents)

        # Calculate similarity matrix
        similarity_matrix = calculate_similarity(tfidf_matrix)

        # Display threshold slider
        similarity_threshold = st.slider('Similarity Threshold', 0.0, 1.0, 0.5, 0.01)

        # Display results
        result_data = []

        for i in range(len(input_documents)):
            doc_results = [
                {
                    'file_name': file_names[j],
                    'similarity': similarity_matrix[i][j]
                }
                for j in range(len(input_documents)) if i != j
            ]

            # Calculate average similarity without considering self-similarity
            avg_similarity = sum(similarity_matrix[i][j] for j in range(len(input_documents)) if i != j) / (len(input_documents) - 1)

            # Filter results based on similarity threshold
            filtered_results = [result for result in doc_results if result['similarity'] >= similarity_threshold]

            result_data.append({
                'file_name': file_names[i],
                'results': filtered_results,
                'average_similarity': avg_similarity
            })

        # Display results
        for result in result_data:
            st.write(f"### {result['file_name']}")
            st.table(result['results'])
            st.write(f"Average Similarity: {result['average_similarity']:.4f}")
            st.write('\n')


if __name__ == '__main__':
    main()
